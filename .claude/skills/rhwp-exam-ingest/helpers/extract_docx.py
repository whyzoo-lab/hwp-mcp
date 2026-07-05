#!/usr/bin/env python3
"""rhwp-exam-ingest helper: DOCX → 텍스트 + 임베디드 이미지 추출.

사용법:
    extract_docx.py <input.docx> <out_dir>

출력:
    <out_dir>/text.txt        — 본문 텍스트 (단락 단위 줄바꿈)
    <out_dir>/img/<name>.png  — 임베디드 이미지 (또는 .jpg 등 원본 확장자)
"""

import sys
import os
import zipfile
from pathlib import Path


def main() -> int:
    if len(sys.argv) < 3:
        print("사용법: extract_docx.py <input.docx> <out_dir>", file=sys.stderr)
        return 1

    inp = Path(sys.argv[1])
    out = Path(sys.argv[2])

    if not inp.exists():
        print(f"오류: 입력 파일이 없습니다: {inp}", file=sys.stderr)
        return 1

    out.mkdir(parents=True, exist_ok=True)
    img_dir = out / "img"
    img_dir.mkdir(exist_ok=True)

    # python-docx 우선 시도 (정밀 텍스트 추출)
    try:
        from docx import Document  # type: ignore
        d = Document(str(inp))
        with open(out / "text.txt", "w", encoding="utf-8") as f:
            for para in d.paragraphs:
                f.write(para.text + "\n")
        print(f"텍스트 추출 ({len(d.paragraphs)} 단락): {out}/text.txt")
    except ImportError:
        # python-docx 없으면 zip 직접 파싱 (lxml 없이 정규식으로 단순 추출)
        import re
        with zipfile.ZipFile(inp, "r") as z:
            with z.open("word/document.xml") as f:
                xml = f.read().decode("utf-8", errors="ignore")
        # <w:t>...</w:t> 텍스트만 추출
        texts = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)
        with open(out / "text.txt", "w", encoding="utf-8") as f:
            f.write("\n".join(texts))
        print(f"텍스트 추출 (정규식 fallback, {len(texts)} 토큰): {out}/text.txt")

    # 임베디드 이미지 추출 (DOCX는 ZIP이므로 word/media/ 폴더에 들어있음)
    img_count = 0
    with zipfile.ZipFile(inp, "r") as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                base = os.path.basename(name)
                if not base:
                    continue
                with z.open(name) as src, open(img_dir / base, "wb") as dst:
                    dst.write(src.read())
                img_count += 1

    print(f"이미지 추출 ({img_count} 개): {img_dir}/")
    return 0


if __name__ == "__main__":
    sys.exit(main())
